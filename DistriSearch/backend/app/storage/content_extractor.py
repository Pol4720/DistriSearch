"""
Content Extractor for DistriSearch.

Extracts text content from various file formats:
- PDF
- Word documents (doc, docx)
- Text files
- HTML
- And more
"""

import asyncio
import io
import logging
import re
from dataclasses import dataclass, field
from typing import Optional, Dict, Any, List
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class ExtractedContent:
    """
    Content extracted from a file.
    
    Attributes:
        text: Extracted text content
        title: Document title (if available)
        metadata: Extracted metadata
        pages: Number of pages (for documents)
        language: Detected language
        word_count: Number of words
        extraction_time_ms: Time taken to extract
    """
    text: str
    title: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    pages: int = 0
    language: Optional[str] = None
    word_count: int = 0
    extraction_time_ms: float = 0.0
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "text": self.text,
            "title": self.title,
            "metadata": self.metadata,
            "pages": self.pages,
            "language": self.language,
            "word_count": self.word_count,
            "extraction_time_ms": self.extraction_time_ms,
        }


class ContentExtractor:
    """
    Extract text content from various file formats.
    
    Supported formats:
    - PDF: Using PyMuPDF (fitz) or pdfplumber
    - Word: Using python-docx
    - Text: Plain text files
    - HTML: Using BeautifulSoup
    - Markdown: Plain text extraction
    """
    
    def __init__(self):
        """Initialize content extractor."""
        self._extractors = {
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
            ".doc": self._extract_doc,
            ".txt": self._extract_text,
            ".html": self._extract_html,
            ".htm": self._extract_html,
            ".md": self._extract_markdown,
            ".csv": self._extract_csv,
            ".json": self._extract_json,
            ".xml": self._extract_xml,
            ".rtf": self._extract_rtf,
        }
        
        logger.info("ContentExtractor initialized")
    
    async def extract(
        self,
        file_data: bytes,
        filename: str,
        content_type: str,
    ) -> ExtractedContent:
        """
        Extract content from a file.
        
        Args:
            file_data: File content bytes
            filename: Original filename
            content_type: MIME type
            
        Returns:
            ExtractedContent object
        """
        start_time = datetime.now()
        
        # Determine extractor based on extension
        extension = Path(filename).suffix.lower()
        extractor = self._extractors.get(extension)
        
        if extractor is None:
            # Try to determine from content type
            extension = self._extension_from_mime(content_type)
            extractor = self._extractors.get(extension)
        
        if extractor is None:
            raise ValueError(f"No extractor for {extension} ({content_type})")
        
        try:
            # Run extraction in thread pool for blocking operations
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                None,
                extractor,
                file_data,
                filename,
            )
            
            extraction_time = (datetime.now() - start_time).total_seconds() * 1000
            result.extraction_time_ms = extraction_time
            result.word_count = len(result.text.split())
            
            logger.debug(
                f"Extracted {result.word_count} words from {filename} "
                f"in {extraction_time:.1f}ms"
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Extraction error for {filename}: {e}")
            raise
    
    def _extension_from_mime(self, content_type: str) -> str:
        """Get file extension from MIME type."""
        mime_map = {
            "application/pdf": ".pdf",
            "application/msword": ".doc",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
            "text/plain": ".txt",
            "text/html": ".html",
            "text/markdown": ".md",
            "text/csv": ".csv",
            "application/json": ".json",
            "application/xml": ".xml",
            "text/xml": ".xml",
        }
        return mime_map.get(content_type, "")
    
    def _extract_pdf(self, file_data: bytes, filename: str) -> ExtractedContent:
        """Extract text from PDF."""
        try:
            # Try PyMuPDF (fitz) first
            import fitz
            
            text_parts = []
            metadata = {}
            pages = 0
            title = None
            
            with fitz.open(stream=file_data, filetype="pdf") as doc:
                pages = len(doc)
                metadata = dict(doc.metadata) if doc.metadata else {}
                title = metadata.get("title")
                
                for page in doc:
                    text_parts.append(page.get_text())
            
            text = "\n".join(text_parts)
            
            return ExtractedContent(
                text=text,
                title=title,
                metadata=metadata,
                pages=pages,
            )
            
        except ImportError:
            # Fallback to pdfplumber
            try:
                import pdfplumber
                
                text_parts = []
                pages = 0
                
                with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                    pages = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                
                text = "\n".join(text_parts)
                
                return ExtractedContent(
                    text=text,
                    pages=pages,
                )
                
            except ImportError:
                raise ImportError(
                    "No PDF library available. Install pymupdf or pdfplumber."
                )
    
    def _extract_docx(self, file_data: bytes, filename: str) -> ExtractedContent:
        """Extract text from DOCX."""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_data))
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            
            # Get title from first heading or paragraph
            title = None
            for para in doc.paragraphs:
                if para.text.strip():
                    title = para.text[:100]
                    break
            
            return ExtractedContent(
                text=text,
                title=title,
            )
            
        except ImportError:
            raise ImportError(
                "python-docx not available. Install with: pip install python-docx"
            )
    
    def _extract_doc(self, file_data: bytes, filename: str) -> ExtractedContent:
        """Extract text from DOC (older Word format)."""
        # For older .doc files, try antiword or textract
        try:
            import subprocess
            import tempfile
            
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as f:
                f.write(file_data)
                temp_path = f.name
            
            try:
                # Try antiword
                result = subprocess.run(
                    ['antiword', temp_path],
                    capture_output=True,
                    text=True,
                )
                
                if result.returncode == 0:
                    return ExtractedContent(text=result.stdout)
                    
            finally:
                import os
                os.unlink(temp_path)
                
        except FileNotFoundError:
            pass
        
        # Fallback - treat as binary with text extraction
        return self._extract_text_from_binary(file_data, filename)
    
    def _extract_text(self, file_data: bytes, filename: str) -> ExtractedContent:
        """Extract text from plain text file."""
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                text = file_data.decode(encoding)
                return ExtractedContent(text=text)
            except UnicodeDecodeError:
                continue
        
        # Last resort - ignore errors
        text = file_data.decode('utf-8', errors='ignore')
        return ExtractedContent(text=text)
    
    def _extract_html(self, file_data: bytes, filename: str) -> ExtractedContent:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup
            
            # Decode HTML
            html = file_data.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html, 'html.parser')
            
            # Extract title
            title = None
            title_tag = soup.find('title')
            if title_tag:
                title = title_tag.get_text().strip()
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Get text
            text = soup.get_text(separator='\n')
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            text = '\n'.join(line for line in lines if line)
            
            # Extract metadata
            metadata = {}
            for meta in soup.find_all('meta'):
                name = meta.get('name', meta.get('property', ''))
                content = meta.get('content', '')
                if name and content:
                    metadata[name] = content
            
            return ExtractedContent(
                text=text,
                title=title,
                metadata=metadata,
            )
            
        except ImportError:
            # Fallback - basic regex extraction
            html = file_data.decode('utf-8', errors='ignore')
            
            # Remove tags
            text = re.sub(r'<[^>]+>', ' ', html)
            text = re.sub(r'\s+', ' ', text).strip()
            
            return ExtractedContent(text=text)
    
    def _extract_markdown(self, file_data: bytes, filename: str) -> ExtractedContent:
        """Extract text from Markdown."""
        text = file_data.decode('utf-8', errors='ignore')
        
        # Extract title from first heading
        title = None
        title_match = re.search(r'^#\s+(.+)$', text, re.MULTILINE)
        if title_match:
            title = title_match.group(1).strip()
        
        # Remove markdown formatting for cleaner text
        # Remove code blocks
        text = re.sub(r'```[\s\S]*?```', '', text)
        text = re.sub(r'`[^`]+`', '', text)
        
        # Remove links but keep text
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        
        # Remove images
        text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', '', text)
        
        # Remove headers markers
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        
        # Clean up whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return ExtractedContent(
            text=text,
            title=title,
        )
    
    def _extract_csv(self, file_data: bytes, filename: str) -> ExtractedContent:
        """Extract text from CSV."""
        import csv
        
        text = file_data.decode('utf-8', errors='ignore')
        
        # Parse CSV and convert to text
        lines = []
        reader = csv.reader(io.StringIO(text))
        
        for row in reader:
            lines.append(' | '.join(row))
        
        return ExtractedContent(text='\n'.join(lines))
    
    def _extract_json(self, file_data: bytes, filename: str) -> ExtractedContent:
        """Extract text from JSON."""
        import json
        
        text = file_data.decode('utf-8', errors='ignore')
        
        try:
            data = json.loads(text)
            
            # Convert JSON to readable text
            text_parts = self._json_to_text(data)
            text = '\n'.join(text_parts)
            
            return ExtractedContent(text=text)
            
        except json.JSONDecodeError:
            return ExtractedContent(text=text)
    
    def _json_to_text(self, obj, prefix: str = "") -> List[str]:
        """Convert JSON object to text lines."""
        lines = []
        
        if isinstance(obj, dict):
            for key, value in obj.items():
                if isinstance(value, (dict, list)):
                    lines.extend(self._json_to_text(value, f"{prefix}{key}."))
                else:
                    lines.append(f"{prefix}{key}: {value}")
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                lines.extend(self._json_to_text(item, f"{prefix}[{i}]."))
        else:
            lines.append(f"{prefix}{obj}")
        
        return lines
    
    def _extract_xml(self, file_data: bytes, filename: str) -> ExtractedContent:
        """Extract text from XML."""
        try:
            from xml.etree import ElementTree
            
            root = ElementTree.fromstring(file_data)
            
            # Extract all text content
            text_parts = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    text_parts.append(elem.text.strip())
                if elem.tail and elem.tail.strip():
                    text_parts.append(elem.tail.strip())
            
            return ExtractedContent(text='\n'.join(text_parts))
            
        except Exception:
            # Fallback - treat as HTML
            return self._extract_html(file_data, filename)
    
    def _extract_rtf(self, file_data: bytes, filename: str) -> ExtractedContent:
        """Extract text from RTF."""
        try:
            from striprtf.striprtf import rtf_to_text
            
            text = rtf_to_text(file_data.decode('utf-8', errors='ignore'))
            return ExtractedContent(text=text)
            
        except ImportError:
            # Basic RTF stripping
            text = file_data.decode('utf-8', errors='ignore')
            
            # Remove RTF commands
            text = re.sub(r'\\[a-z]+\d*\s?', '', text)
            text = re.sub(r'\{|\}', '', text)
            
            return ExtractedContent(text=text)
    
    def _extract_text_from_binary(
        self,
        file_data: bytes,
        filename: str,
    ) -> ExtractedContent:
        """Try to extract text from binary file."""
        # Look for ASCII text sequences
        text_parts = []
        current_text = []
        
        for byte in file_data:
            if 32 <= byte <= 126 or byte in (9, 10, 13):  # Printable ASCII
                current_text.append(chr(byte))
            else:
                if len(current_text) > 10:  # Minimum text length
                    text_parts.append(''.join(current_text))
                current_text = []
        
        if len(current_text) > 10:
            text_parts.append(''.join(current_text))
        
        return ExtractedContent(text='\n'.join(text_parts))
    
    @staticmethod
    def detect_language(text: str) -> str:
        """
        Detect language of text.
        
        Uses simple heuristics for common languages.
        """
        # Common words for language detection
        language_indicators = {
            'en': ['the', 'and', 'is', 'are', 'was', 'were', 'have', 'has'],
            'es': ['el', 'la', 'los', 'las', 'de', 'en', 'que', 'es'],
            'fr': ['le', 'la', 'les', 'de', 'des', 'est', 'sont', 'dans'],
            'de': ['der', 'die', 'das', 'und', 'ist', 'sind', 'ein', 'eine'],
            'pt': ['o', 'a', 'os', 'as', 'de', 'da', 'do', 'que'],
        }
        
        words = text.lower().split()
        word_set = set(words)
        
        scores = {}
        for lang, indicators in language_indicators.items():
            score = sum(1 for word in indicators if word in word_set)
            scores[lang] = score
        
        if scores:
            best_lang = max(scores, key=scores.get)
            if scores[best_lang] > 0:
                return best_lang
        
        return 'unknown'
